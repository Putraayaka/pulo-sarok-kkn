import google.generativeai as genai
import hashlib
import qrcode
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from django.conf import settings
from django.core.files.base import ContentFile
from django.utils import timezone
from datetime import datetime, timedelta
import re
import logging
import os
import json
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import LetterSettings, APIKeySettings, Letter, LetterAIValidation

logger = logging.getLogger(__name__)

class GeminiAIService:
    """Service untuk integrasi dengan Google Gemini AI"""
    
    def __init__(self):
        self.api_key = self._get_api_key()
        if self.api_key:
            genai.configure(api_key=self.api_key)
            self.model = genai.GenerativeModel('gemini-2.5-pro')
        else:
            self.model = None
            logger.warning("Gemini API key tidak ditemukan")
    
    def _get_api_key(self):
        """Ambil API key dari database"""
        try:
            api_settings = APIKeySettings.objects.first()
            return api_settings.gemini_api_key if api_settings else None
        except Exception as e:
            logger.error(f"Error mengambil API key: {e}")
            return None
    
    def validate_letter(self, content, letter_type=None):
        """Validasi konten surat menggunakan AI"""
        if not self.model:
            return {
                'is_valid': False,
                'score': 0,
                'suggestions': ['API key Gemini tidak tersedia'],
                'errors': ['Konfigurasi AI tidak lengkap']
            }
        
        try:
            prompt = f"""
            Analisis konten surat berikut dan berikan validasi:
            
            Jenis Surat: {letter_type or 'Tidak ditentukan'}
            Konten: {content}
            
            Berikan analisis dalam format JSON dengan struktur:
            {{
                "is_valid": boolean,
                "score": number (0-100),
                "suggestions": ["saran1", "saran2"],
                "errors": ["error1", "error2"],
                "grammar_check": "hasil pengecekan tata bahasa",
                "formality_check": "hasil pengecekan formalitas",
                "completeness_check": "hasil pengecekan kelengkapan"
            }}
            
            Fokus pada:
            1. Tata bahasa Indonesia yang benar
            2. Formalitas bahasa surat resmi
            3. Kelengkapan struktur surat
            4. Kejelasan maksud dan tujuan
            """
            
            response = self.model.generate_content(prompt)
            result = self._parse_ai_response(response.text)
            return result
            
        except Exception as e:
            logger.error(f"Error validasi AI: {e}")
            return {
                'is_valid': False,
                'score': 0,
                'suggestions': [f'Error AI: {str(e)}'],
                'errors': ['Gagal melakukan validasi AI']
            }
    
    def improve_letter(self, content, letter_type=None):
        """Saran perbaikan konten surat"""
        if not self.model:
            return {'improved_content': content, 'suggestions': ['AI tidak tersedia']}
        
        try:
            prompt = f"""
            Perbaiki dan tingkatkan kualitas surat berikut:
            
            Jenis Surat: {letter_type or 'Tidak ditentukan'}
            Konten Asli: {content}
            
            Berikan hasil dalam format JSON:
            {{
                "improved_content": "konten yang diperbaiki",
                "suggestions": ["saran perbaikan"],
                "changes_made": ["perubahan yang dilakukan"]
            }}
            
            Fokus perbaikan:
            1. Tata bahasa dan ejaan
            2. Struktur kalimat
            3. Formalitas bahasa
            4. Kejelasan komunikasi
            """
            
            response = self.model.generate_content(prompt)
            return self._parse_ai_response(response.text)
            
        except Exception as e:
            logger.error(f"Error improve AI: {e}")
            return {'improved_content': content, 'suggestions': [f'Error: {str(e)}']}
    
    def generate_letter(self, letter_type, purpose, recipient, additional_info=""):
        """Generate konten surat baru menggunakan AI"""
        if not self.model:
            return {'content': '', 'suggestions': ['AI tidak tersedia']}
        
        try:
            prompt = f"""
            Buatkan draft surat resmi dengan detail berikut:
            
            Jenis Surat: {letter_type}
            Tujuan: {purpose}
            Penerima: {recipient}
            Informasi Tambahan: {additional_info}
            
            Berikan hasil dalam format JSON:
            {{
                "content": "konten surat lengkap",
                "subject": "perihal surat",
                "suggestions": ["saran tambahan"]
            }}
            
            Struktur surat harus mencakup:
            1. Kop surat (placeholder)
            2. Nomor, tanggal, perihal
            3. Alamat tujuan
            4. Pembukaan
            5. Isi surat
            6. Penutup
            7. Tanda tangan (placeholder)
            
            Gunakan bahasa Indonesia formal dan resmi.
            """
            
            response = self.model.generate_content(prompt)
            return self._parse_ai_response(response.text)
            
        except Exception as e:
            logger.error(f"Error generate AI: {e}")
            return {'content': '', 'suggestions': [f'Error: {str(e)}']}
    
    def summarize_letter(self, content):
        """Buat ringkasan surat"""
        if not self.model:
            return {'summary': 'AI tidak tersedia', 'key_points': []}
        
        try:
            prompt = f"""
            Buatkan ringkasan dari surat berikut:
            
            Konten: {content}
            
            Berikan hasil dalam format JSON:
            {{
                "summary": "ringkasan singkat",
                "key_points": ["poin penting 1", "poin penting 2"],
                "word_count": jumlah_kata,
                "reading_time": "estimasi waktu baca dalam menit"
            }}
            """
            
            response = self.model.generate_content(prompt)
            return self._parse_ai_response(response.text)
            
        except Exception as e:
            logger.error(f"Error summarize AI: {e}")
            return {'summary': f'Error: {str(e)}', 'key_points': []}
    
    def _parse_ai_response(self, response_text):
        """Parse response AI menjadi JSON"""
        try:
            # Coba ekstrak JSON dari response
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            if json_start != -1 and json_end != -1:
                json_str = response_text[json_start:json_end]
                return json.loads(json_str)
            else:
                # Fallback jika tidak ada JSON
                return {'content': response_text, 'suggestions': []}
        except Exception as e:
            logger.error(f"Error parsing AI response: {e}")
            return {'content': response_text, 'suggestions': []}

class LetterValidationService:
    """Service untuk validasi surat"""
    
    @staticmethod
    def validate_content(content):
        """Validasi dasar konten surat"""
        errors = []
        warnings = []
        
        if not content or len(content.strip()) < 10:
            errors.append("Konten surat terlalu pendek")
        
        if len(content) > 10000:
            warnings.append("Konten surat sangat panjang")
        
        # Cek struktur dasar
        if 'Kepada' not in content and 'Yth.' not in content:
            warnings.append("Alamat tujuan tidak ditemukan")
        
        if 'Hormat kami' not in content and 'Wassalam' not in content:
            warnings.append("Penutup surat tidak ditemukan")
        
        return {
            'is_valid': len(errors) == 0,
            'errors': errors,
            'warnings': warnings
        }
    
    @staticmethod
    def check_required_fields(letter_data):
        """Cek field yang wajib diisi"""
        required_fields = ['subject', 'content', 'letter_type']
        missing_fields = []
        
        for field in required_fields:
            if not letter_data.get(field):
                missing_fields.append(field)
        
        return {
            'is_complete': len(missing_fields) == 0,
            'missing_fields': missing_fields
        }

class LetterNumberingService:
    """Service untuk penomoran surat otomatis"""
    
    @staticmethod
    def generate_letter_number(letter_type, date=None):
        """Generate nomor surat otomatis"""
        if not date:
            date = timezone.now().date()
        
        try:
            settings_obj = LetterSettings.objects.first()
            if not settings_obj:
                # Default format jika belum ada settings
                format_template = "{sequence:03d}/{type_code}/{month:02d}/{year}"
                type_code = "UMUM"
            else:
                format_template = settings_obj.letter_number_format
                type_code = letter_type.code if hasattr(letter_type, 'code') else "UMUM"
            
            # Hitung sequence number untuk bulan ini
            current_month = date.month
            current_year = date.year
            
            last_letter = Letter.objects.filter(
                created_at__year=current_year,
                created_at__month=current_month
            ).order_by('-id').first()
            
            sequence = 1
            if last_letter and last_letter.letter_number:
                # Extract sequence dari nomor terakhir
                try:
                    parts = last_letter.letter_number.split('/')
                    if parts:
                        sequence = int(parts[0]) + 1
                except:
                    sequence = 1
            
            # Format nomor surat
            letter_number = format_template.format(
                sequence=sequence,
                type_code=type_code,
                month=current_month,
                year=current_year,
                day=date.day
            )
            
            return letter_number
            
        except Exception as e:
            logger.error(f"Error generating letter number: {e}")
            # Fallback format
            return f"{date.year}{date.month:02d}{date.day:02d}-001"

class LetterExportService:
    """Service untuk export surat ke berbagai format"""
    
    @staticmethod
    def export_to_pdf(letter):
        """Export surat ke PDF"""
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Kop surat
            settings_obj = LetterSettings.objects.first()
            if settings_obj:
                # Header dengan kop surat
                header_style = ParagraphStyle(
                    'CustomHeader',
                    parent=styles['Heading1'],
                    alignment=TA_CENTER,
                    fontSize=14,
                    spaceAfter=20
                )
                
                story.append(Paragraph(settings_obj.village_name, header_style))
                story.append(Paragraph(settings_obj.village_address, styles['Normal']))
                story.append(Spacer(1, 20))
            
            # Nomor dan tanggal surat
            info_style = ParagraphStyle(
                'InfoStyle',
                parent=styles['Normal'],
                alignment=TA_LEFT,
                fontSize=10
            )
            
            story.append(Paragraph(f"Nomor: {letter.letter_number or '-'}", info_style))
            story.append(Paragraph(f"Tanggal: {letter.created_at.strftime('%d %B %Y')}", info_style))
            story.append(Paragraph(f"Perihal: {letter.subject}", info_style))
            story.append(Spacer(1, 20))
            
            # Konten surat
            content_style = ParagraphStyle(
                'ContentStyle',
                parent=styles['Normal'],
                alignment=TA_LEFT,
                fontSize=11,
                leading=14
            )
            
            # Split konten menjadi paragraf
            paragraphs = letter.content.split('\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para, content_style))
                    story.append(Spacer(1, 6))
            
            # Tanda tangan
            if settings_obj and settings_obj.village_head_name:
                story.append(Spacer(1, 30))
                signature_style = ParagraphStyle(
                    'SignatureStyle',
                    parent=styles['Normal'],
                    alignment=TA_RIGHT,
                    fontSize=11
                )
                story.append(Paragraph(f"Kepala Desa,", signature_style))
                story.append(Spacer(1, 40))
                story.append(Paragraph(settings_obj.village_head_name, signature_style))
            
            doc.build(story)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            return None
    
    @staticmethod
    def export_to_docx(letter):
        """Export surat ke DOCX"""
        try:
            doc = Document()
            
            # Kop surat
            settings_obj = LetterSettings.objects.first()
            if settings_obj:
                header = doc.sections[0].header
                header_para = header.paragraphs[0]
                header_para.text = settings_obj.village_name
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Alamat desa
                header_para2 = header.add_paragraph(settings_obj.village_address)
                header_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Informasi surat
            doc.add_paragraph(f"Nomor: {letter.letter_number or '-'}")
            doc.add_paragraph(f"Tanggal: {letter.created_at.strftime('%d %B %Y')}")
            doc.add_paragraph(f"Perihal: {letter.subject}")
            doc.add_paragraph()  # Empty line
            
            # Konten surat
            paragraphs = letter.content.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)
            
            # Tanda tangan
            if settings_obj and settings_obj.village_head_name:
                doc.add_paragraph()  # Empty line
                signature_para = doc.add_paragraph("Kepala Desa,")
                signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                doc.add_paragraph()  # Space for signature
                doc.add_paragraph()  # Space for signature
                
                name_para = doc.add_paragraph(settings_obj.village_head_name)
                name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            return None

# Utility functions
def generate_qr_code(data, size=(200, 200)):
    """Generate QR code untuk surat"""
    try:
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(data)
        qr.make(fit=True)
        
        img = qr.make_image(fill_color="black", back_color="white")
        img = img.resize(size)
        
        buffer = BytesIO()
        img.save(buffer, format='PNG')
        buffer.seek(0)
        
        return ContentFile(buffer.getvalue(), name='qr_code.png')
    except Exception as e:
        logger.error(f"Error generating QR code: {e}")
        return None

def create_letterhead_image(village_name, village_address, logo_path=None):
    """Buat gambar kop surat"""
    try:
        # Buat image kosong
        img = Image.new('RGB', (800, 200), color='white')
        draw = ImageDraw.Draw(img)
        
        # Font (gunakan font default jika tidak ada)
        try:
            title_font = ImageFont.truetype("arial.ttf", 24)
            subtitle_font = ImageFont.truetype("arial.ttf", 16)
        except:
            title_font = ImageFont.load_default()
            subtitle_font = ImageFont.load_default()
        
        # Tulis nama desa
        title_bbox = draw.textbbox((0, 0), village_name, font=title_font)
        title_width = title_bbox[2] - title_bbox[0]
        title_x = (800 - title_width) // 2
        draw.text((title_x, 50), village_name, fill='black', font=title_font)
        
        # Tulis alamat
        addr_bbox = draw.textbbox((0, 0), village_address, font=subtitle_font)
        addr_width = addr_bbox[2] - addr_bbox[0]
        addr_x = (800 - addr_width) // 2
        draw.text((addr_x, 90), village_address, fill='black', font=subtitle_font)
        
        # Garis bawah
        draw.line([(50, 150), (750, 150)], fill='black', width=2)
        
        buffer = BytesIO()
        img.save(buffer, format='PNG')
        buffer.seek(0)
        
        return ContentFile(buffer.getvalue(), name='letterhead.png')
    except Exception as e:
        logger.error(f"Error creating letterhead: {e}")
        return None

def calculate_reading_time(content):
    """Hitung estimasi waktu baca"""
    words = len(content.split())
    # Asumsi 200 kata per menit
    minutes = max(1, words // 200)
    return minutes

def extract_text_statistics(content):
    """Ekstrak statistik teks"""
    words = len(content.split())
    characters = len(content)
    characters_no_spaces = len(content.replace(' ', ''))
    paragraphs = len([p for p in content.split('\n') if p.strip()])
    
    return {
        'words': words,
        'characters': characters,
        'characters_no_spaces': characters_no_spaces,
        'paragraphs': paragraphs,
        'reading_time': calculate_reading_time(content)
    }

def generate_letter_hash(content, letter_number):
    """Generate hash unik untuk surat"""
    data = f"{content}{letter_number}{timezone.now().isoformat()}"
    return hashlib.sha256(data.encode()).hexdigest()[:16]

def validate_file_upload(file):
    """Validasi file upload"""
    max_size = 5 * 1024 * 1024  # 5MB
    allowed_types = ['pdf', 'doc', 'docx', 'jpg', 'jpeg', 'png']
    
    if file.size > max_size:
        return False, "File terlalu besar (maksimal 5MB)"
    
    file_ext = file.name.split('.')[-1].lower()
    if file_ext not in allowed_types:
        return False, f"Tipe file tidak diizinkan. Gunakan: {', '.join(allowed_types)}"
    
    return True, "File valid"
import requests
from django.conf import settings
from django.utils import timezone
from .models import APIKeySettings, LetterAIValidation, Letter
import logging

logger = logging.getLogger(__name__)

class GeminiAIService:
    """Service for integrating with Google Gemini AI"""
    
    def __init__(self):
        self.api_key = self.get_api_key()
        self.base_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-pro:generateContent"
    
    def get_api_key(self):
        """Get Gemini API key from database"""
        try:
            api_settings = APIKeySettings.objects.filter(
                service_name='gemini',
                is_active=True
            ).first()
            
            if api_settings and api_settings.is_valid():
                return api_settings.get_api_key()
            return None
        except Exception as e:
            logger.error(f"Error getting API key: {e}")
            return None
    
    def validate_letter_content(self, letter_content, letter_type=None):
        """Validate letter content using Gemini AI"""
        if not self.api_key:
            return {
                'status': 'error',
                'message': 'API key tidak tersedia',
                'confidence_score': 0.0
            }
        
        try:
            # Prepare prompt for validation
            prompt = self._create_validation_prompt(letter_content, letter_type)
            
            # Make API request
            response = self._make_api_request(prompt)
            
            if response:
                # Increment API usage
                self._increment_api_usage()
                return self._parse_validation_response(response)
            
            return {
                'status': 'error',
                'message': 'Gagal mendapatkan respons dari AI',
                'confidence_score': 0.0
            }
            
        except Exception as e:
            logger.error(f"Error validating letter: {e}")
            return {
                'status': 'error',
                'message': f'Error: {str(e)}',
                'confidence_score': 0.0
            }
    
    def generate_letter_suggestions(self, letter_content, letter_type=None):
        """Generate suggestions for improving letter content"""
        if not self.api_key:
            return []
        
        try:
            prompt = self._create_suggestion_prompt(letter_content, letter_type)
            response = self._make_api_request(prompt)
            
            if response:
                self._increment_api_usage()
                return self._parse_suggestions_response(response)
            
            return []
            
        except Exception as e:
            logger.error(f"Error generating suggestions: {e}")
            return []
    
    def generate_letter_template(self, letter_type, purpose, recipient_type=None):
        """Generate letter template based on type and purpose"""
        if not self.api_key:
            return None
        
        try:
            prompt = self._create_template_prompt(letter_type, purpose, recipient_type)
            response = self._make_api_request(prompt)
            
            if response:
                self._increment_api_usage()
                return self._parse_template_response(response)
            
            return None
            
        except Exception as e:
            logger.error(f"Error generating template: {e}")
            return None
    
    def _create_validation_prompt(self, content, letter_type):
        """Create prompt for letter validation"""
        return f"""
Anda adalah asisten AI yang ahli dalam validasi surat resmi pemerintahan desa di Indonesia.

Tugas: Validasi konten surat berikut dan berikan penilaian.

Jenis Surat: {letter_type or 'Umum'}
Konten Surat:
{content}

Harap analisis:
1. Struktur dan format surat
2. Penggunaan bahasa formal
3. Kelengkapan informasi
4. Kesesuaian dengan standar surat dinas
5. Tata bahasa dan ejaan

Berikan respons dalam format JSON:
{{
    "is_valid": true/false,
    "confidence_score": 0.0-1.0,
    "issues": ["daftar masalah yang ditemukan"],
    "suggestions": ["saran perbaikan"],
    "overall_assessment": "penilaian keseluruhan"
}}
"""
    
    def _create_suggestion_prompt(self, content, letter_type):
        """Create prompt for generating suggestions"""
        return f"""
Anda adalah asisten AI yang ahli dalam penulisan surat resmi pemerintahan desa.

Tugas: Berikan saran untuk memperbaiki surat berikut.

Jenis Surat: {letter_type or 'Umum'}
Konten Surat:
{content}

Berikan 5-10 saran konkret untuk memperbaiki surat ini dalam format JSON:
{{
    "suggestions": [
        {{
            "type": "struktur/bahasa/format/konten",
            "description": "deskripsi saran",
            "priority": "high/medium/low"
        }}
    ]
}}
"""
    
    def _create_template_prompt(self, letter_type, purpose, recipient_type):
        """Create prompt for generating letter template"""
        return f"""
Anda adalah asisten AI yang ahli dalam pembuatan template surat resmi pemerintahan desa.

Tugas: Buat template surat dengan spesifikasi berikut:
- Jenis Surat: {letter_type}
- Tujuan: {purpose}
- Penerima: {recipient_type or 'Umum'}

Buat template lengkap dengan:
1. Header/kop surat (gunakan placeholder {{village_name}}, {{village_address}})
2. Nomor surat (gunakan placeholder {{letter_number}})
3. Tanggal (gunakan placeholder {{date}})
4. Perihal
5. Isi surat dengan struktur yang tepat
6. Penutup dan tanda tangan (gunakan placeholder {{head_name}})

Berikan respons dalam format JSON:
{{
    "template_name": "nama template",
    "content": "konten template lengkap",
    "variables": ["daftar variabel yang digunakan"],
    "description": "deskripsi template"
}}
"""
    
    def _make_api_request(self, prompt):
        """Make API request to Gemini"""
        headers = {
            'Content-Type': 'application/json',
        }
        
        data = {
            "contents": [{
                "parts": [{
                    "text": prompt
                }]
            }],
            "generationConfig": {
                "temperature": 0.7,
                "topK": 40,
                "topP": 0.95,
                "maxOutputTokens": 2048,
            }
        }
        
        url = f"{self.base_url}?key={self.api_key}"
        
        response = requests.post(url, headers=headers, json=data, timeout=30)
        
        if response.status_code == 200:
            return response.json()
        else:
            logger.error(f"API request failed: {response.status_code} - {response.text}")
            return None
    
    def _parse_validation_response(self, response):
        """Parse validation response from Gemini"""
        try:
            content = response['candidates'][0]['content']['parts'][0]['text']
            
            # Try to extract JSON from response
            start_idx = content.find('{')
            end_idx = content.rfind('}') + 1
            
            if start_idx != -1 and end_idx != -1:
                json_str = content[start_idx:end_idx]
                result = json.loads(json_str)
                
                return {
                    'status': 'valid' if result.get('is_valid', False) else 'invalid',
                    'confidence_score': result.get('confidence_score', 0.0),
                    'issues': result.get('issues', []),
                    'suggestions': result.get('suggestions', []),
                    'assessment': result.get('overall_assessment', '')
                }
            
            # Fallback if JSON parsing fails
            return {
                'status': 'processed',
                'confidence_score': 0.5,
                'issues': [],
                'suggestions': [content],
                'assessment': content
            }
            
        except Exception as e:
            logger.error(f"Error parsing validation response: {e}")
            return {
                'status': 'error',
                'confidence_score': 0.0,
                'issues': ['Gagal memproses respons AI'],
                'suggestions': [],
                'assessment': 'Error dalam validasi'
            }
    
    def _parse_suggestions_response(self, response):
        """Parse suggestions response from Gemini"""
        try:
            content = response['candidates'][0]['content']['parts'][0]['text']
            
            # Try to extract JSON
            start_idx = content.find('{')
            end_idx = content.rfind('}') + 1
            
            if start_idx != -1 and end_idx != -1:
                json_str = content[start_idx:end_idx]
                result = json.loads(json_str)
                return result.get('suggestions', [])
            
            # Fallback: split by lines
            lines = content.split('\n')
            suggestions = []
            for line in lines:
                line = line.strip()
                if line and not line.startswith('#'):
                    suggestions.append({
                        'type': 'general',
                        'description': line,
                        'priority': 'medium'
                    })
            
            return suggestions[:10]  # Limit to 10 suggestions
            
        except Exception as e:
            logger.error(f"Error parsing suggestions response: {e}")
            return []
    
    def _parse_template_response(self, response):
        """Parse template response from Gemini"""
        try:
            content = response['candidates'][0]['content']['parts'][0]['text']
            
            # Try to extract JSON
            start_idx = content.find('{')
            end_idx = content.rfind('}') + 1
            
            if start_idx != -1 and end_idx != -1:
                json_str = content[start_idx:end_idx]
                result = json.loads(json_str)
                return result
            
            # Fallback: return content as template
            return {
                'template_name': 'Template Generated',
                'content': content,
                'variables': [],
                'description': 'Template yang dihasilkan AI'
            }
            
        except Exception as e:
            logger.error(f"Error parsing template response: {e}")
            return None
    
    def _increment_api_usage(self):
        """Increment API usage counter"""
        try:
            api_settings = APIKeySettings.objects.filter(
                service_name='gemini',
                is_active=True
            ).first()
            
            if api_settings:
                api_settings.increment_usage()
                
        except Exception as e:
            logger.error(f"Error incrementing API usage: {e}")


class LetterValidationService:
    """Service for validating letters using AI and business rules"""
    
    def __init__(self):
        self.gemini_service = GeminiAIService()
    
    def validate_letter(self, letter):
        """Validate a letter using AI and business rules"""
        try:
            # Check if letter requires AI validation
            if not letter.requires_ai_validation:
                return self._create_validation_result(letter, 'skipped', 1.0, [])
            
            # Perform AI validation
            ai_result = self.gemini_service.validate_letter_content(
                letter.content,
                letter.letter_type.name if letter.letter_type else None
            )
            
            # Create or update validation record
            validation, created = LetterAIValidation.objects.get_or_create(
                letter=letter,
                defaults={
                    'status': ai_result['status'],
                    'confidence_score': ai_result['confidence_score'],
                    'suggestions': ai_result.get('suggestions', []),
                    'validation_details': ai_result,
                    'validated_at': timezone.now()
                }
            )
            
            if not created:
                validation.status = ai_result['status']
                validation.confidence_score = ai_result['confidence_score']
                validation.suggestions = ai_result.get('suggestions', [])
                validation.validation_details = ai_result
                validation.validated_at = timezone.now()
                validation.save()
            
            return validation
            
        except Exception as e:
            logger.error(f"Error validating letter {letter.id}: {e}")
            return self._create_validation_result(letter, 'error', 0.0, [str(e)])
    
    def _create_validation_result(self, letter, status, confidence, suggestions):
        """Create validation result"""
        validation, created = LetterAIValidation.objects.get_or_create(
            letter=letter,
            defaults={
                'status': status,
                'confidence_score': confidence,
                'suggestions': suggestions,
                'validated_at': timezone.now()
            }
        )
        return validation


class LetterNumberingService:
    """Service for generating letter numbers"""
    
    @staticmethod
    def generate_letter_number(letter_type_code, settings=None):
        """Generate letter number based on settings"""
        from .models import LetterSettings
        from datetime import datetime
        
        if not settings:
            settings = LetterSettings.objects.filter(is_active=True).first()
        
        if not settings:
            # Fallback to default format
            now = datetime.now()
            count = Letter.objects.filter(
                letter_type__code=letter_type_code,
                created_at__year=now.year,
                created_at__month=now.month
            ).count() + 1
            return f"{letter_type_code}/{count:03d}/{now.month:02d}/{now.year}"
        
        return settings.get_next_letter_number(letter_type_code)


class LetterExportService:
    """Service for exporting letters to various formats"""
    
    @staticmethod
    def export_to_pdf(letter):
        """Export letter to PDF"""
        try:
            return letter.generate_pdf()
        except Exception as e:
            logger.error(f"Error exporting letter {letter.id} to PDF: {e}")
            return False
    
    @staticmethod
    def export_to_docx(letter):
        """Export letter to DOCX format"""
        try:
            from docx import Document
            from io import BytesIO
            from django.core.files import File
            
            doc = Document()
            
            # Add letter content
            doc.add_heading(letter.subject, 0)
            doc.add_paragraph(f"Nomor: {letter.letter_number or 'Draft'}")
            doc.add_paragraph(f"Tanggal: {letter.created_at.strftime('%d %B %Y')}")
            doc.add_paragraph()
            
            # Add content paragraphs
            for paragraph in letter.content.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Save to BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
            
        except ImportError:
            logger.error("python-docx not installed")
            return None
        except Exception as e:
            logger.error(f"Error exporting letter {letter.id} to DOCX: {e}")
            return None